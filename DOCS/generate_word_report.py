import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Arial'
    return heading

def add_paragraph(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def create_report():
    doc = Document()

    # --- TITLE PAGE ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Full-Stack Engineering\nProject Report\nSemester-VI (Batch-2023)\n\n")
    run.font.size = Pt(16)
    run.bold = True
    
    project_title = doc.add_paragraph()
    project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = project_title.add_run("VOTEPULSE\n(Online Voting System)\n\n")
    run2.font.size = Pt(24)
    run2.bold = True
    
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = details.add_run("Submitted in partial fulfilment of the requirement for the Course\nof\nCOMPUTER SCIENCE AND ENGINEERING\nB.E. Batch-2023\nMAY-2026\n\n")
    run3.font.size = Pt(14)

    sub_details = doc.add_paragraph()
    sub_details.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_details.add_run("Supervised By:\t\t\t\tSubmitted By:\n").bold = True
    sub_details.add_run("Mr. Nurul Islam\t\t\t\tHarshit (2310990684)\n\t\t\t\t\tIshita (2310990695)\n\t\t\t\t\tSomya Sharma (2310991297)\n\t\t\t\t\tLiza Garg (2310991334)\n\n")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Department of Computer Science and Engineering\nChitkara University Institute of Engineering & Technology,\nChitkara University, Punjab").bold = True
    
    doc.add_page_break()

    # --- CANDIDATE'S DECLARATION ---
    add_heading(doc, "CANDIDATE'S DECLARATION", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "\nWe, Harshit (2310990684), Somya Sharma (2310991297), Ishita (2310990695), and Liza Garg (2310991334), students of B.E.-2023, Department of Computer Science and Engineering, Chitkara University, Punjab, hereby declare that the Project Report entitled “VotePulse” is an original work carried out by us under the guidance of Mr. Nurul Islam.\n\nThe information and data presented in this report are authentic to the best of our knowledge. This report has not been submitted to any other institute or university for the award of any other degree, diploma, or course.\n\n\nPlace: Chitkara University, Punjab\nDate: 24 May 2026\n")
    doc.add_page_break()

    # --- ACKNOWLEDGEMENT ---
    add_heading(doc, "ACKNOWLEDGEMENT", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "\nIt is our pleasure to express our profound gratitude and deep regards to our guide Mr. Nurul Islam for his exemplary guidance, monitoring and constant encouragement throughout the course of this project. The blessing, help and guidance given by him time to time shall carry us a long way in the journey of life on which we are about to embark.\n\nWe also take this opportunity to express a deep sense of gratitude to Chitkara University for their cordial support, valuable information and guidance, which helped us in completing this task through various stages.\n")
    doc.add_page_break()

    # --- ABSTRACT ---
    add_heading(doc, "Abstract", 1)
    add_paragraph(doc, "The landscape of democratic participation is undergoing a massive shift towards digital transformation. VotePulse is a secure, efficient, and highly scalable digital platform designed to modernize the traditional election process. It provides voters with the unprecedented convenience of casting their ballots from anywhere in the world, overcoming geographical boundaries, mobility limitations, and significantly reducing the massive logistical costs associated with physical polling stations. At its core, VotePulse ensures data integrity and security through robust authentication mechanisms, strict role-based access control (RBAC), and encrypted data storage. By bridging the gap between technology and civic duty, this project aims to foster higher voter turnout, eliminate human error in vote tallying, and build unwavering trust in the electoral process.")
    doc.add_page_break()

    chapters = [
        ("Chapter 1: Introduction", [
            "Democracy fundamentally relies on the active, unhindered participation of its citizens, and voting is the primary mechanism of this participation. However, traditional paper-based or localized electronic voting machines often present significant barriers to entry for remote voters, individuals with disabilities, and the elderly. Furthermore, organizing physical elections requires massive logistical planning, financial expenditure, and human resources.",
            "This project introduces VotePulse, a modern, full-stack web application that digitizes the entire election workflow from end to end. Whether it is a university student council election, a corporate board decision, or a municipal voting event, VotePulse offers a seamless, secure experience. From user registration and automated email verification to election management, candidate allocation, and real-time result tallying, VotePulse provides an intuitive digital election environment that prioritizes both security and the user experience.",
            "The rapid advancement of web technologies, particularly the MERN/PERN stack paradigms, has enabled the creation of highly responsive single-page applications that mimic the performance of native desktop software. By utilizing these technologies, VotePulse ensures a smooth, uninterrupted user experience even under heavy loads during peak voting hours. Security is baked into the architecture from day one, employing cryptographic hashing for passwords and stateless JSON Web Tokens for session management."
        ]),
        ("Chapter 2: Problem Statement", [
            "The conventional, physical voting system faces a multitude of critical challenges that limit its effectiveness and reach:",
            "1. Accessibility Barriers: Voters must be physically present at specific locations within designated timeframes. This poses a severe challenge for the elderly, individuals with physical disabilities, and those living or working abroad (expatriates). The requirement of physical presence disenfranchises a significant portion of the eligible voting population.",
            "2. Massive Logistical Overhead: Organizing physical elections requires immense manpower. Expenses include printing paper ballots, renting venues, hiring security personnel, and transporting ballot boxes securely. These recurring costs strain the budgets of organizations and municipalities.",
            "3. Inefficiency and Time Consumption: Standing in long queues discourages voter turnout. Furthermore, the manual counting of paper ballots is extremely slow and delays the announcement of results. The longer the delay, the higher the anxiety and potential for disputes.",
            "4. Security Risks and Tampering: Centralized physical boxes and paper ballots are susceptible to damage, loss, ballot stuffing, and human error during the counting process. Transporting physical ballots introduces a massive chain-of-custody vulnerability.",
            "5. Lack of Transparency: Voters have no way to verify that their vote was recorded accurately without compromising the secret ballot principle."
        ]),
        ("Chapter 3: Goals/Objectives & Key Learnings", [
            "Goals & Objectives:",
            "- Uncompromising Security: Implement robust user authentication, mandatory email verification for identity confirmation, and secure password hashing to protect user accounts. Security is not an afterthought; it is the foundation of VotePulse.",
            "- Transparency & Accuracy: Ensure that votes are accurately recorded in the database, immutable once cast, and that election results are tallied and displayed in real-time without manual intervention.",
            "- Intuitive Usability: Provide a highly responsive, modern, and intuitive user interface (UI) for both voters (to effortlessly cast their votes) and administrators (to efficiently manage complex elections).",
            "- Strict Role-Based Access Control (RBAC): Clearly differentiate functionalities, ensuring that system administrators have full control over election management, while standard voters are restricted to viewing and participating in active elections.",
            "\nKey Learnings:",
            "- Modern Authentication: Deep understanding of implementing stateless, token-based authentication using JSON Web Tokens (JWT) and secure password management using Bcrypt.",
            "- Database Design: Handling complex relational data structures, foreign keys, and cascading deletes using an Object-Relational Mapper (Sequelize) with a MySQL database.",
            "- Frontend Architecture: Building dynamic, responsive Single Page Applications (SPAs) using React 19, managing complex application state, and utilizing utility-first CSS frameworks like Tailwind CSS combined with Material-UI components.",
            "- Third-Party Integrations: Successfully integrating external APIs and services, such as Cloudinary for seamless image uploads (candidate profiles) and Nodemailer with SMTP servers for reliable transactional email delivery."
        ]),
        ("Chapter 4: Functional & Non-Functional Requirements", [
            "Functional Requirements:",
            "- User Registration: Users can sign up providing their name, email, and password.",
            "- Email Verification: The system sends a verification code/link to the registered email to confirm identity before allowing login.",
            "- User Login: Secure authentication returning a JWT session token.",
            "- Role-based Dashboard: Displays specific options based on whether the user is an 'admin' or 'voter'.",
            "- Election Management: Add, modify, or close electoral events (title, start/end dates).",
            "- Candidate Management: Add candidates to elections, upload their profile pictures, and assign parties.",
            "- View Active Elections: Users can browse ongoing elections and view the list of participating candidates.",
            "- Cast Vote: Users can securely cast a single, immutable vote per active election.",
            "- Result Tallying: System automatically aggregates votes and declares real-time tallies/winners.",
            "- Manage Users: Administrators can view all registered users and their verification statuses.",
            "\nNon-Functional Requirements:",
            "- Security (Auth): Passwords must be hashed using Bcrypt (salt rounds >= 10). Endpoints must use JWT authorization headers.",
            "- Security (Integrity): Database constraints must strictly prevent duplicate votes from the same user ID in a single election. A unique composite index handles this at the database level.",
            "- Performance: API endpoints (except heavy image uploads) should respond in < 200ms.",
            "- Scalability: Architecture must handle a high volume of concurrent POST requests during peak voting hours.",
            "- Responsiveness: The UI must adapt seamlessly across mobile phones, tablets, and desktop displays (Mobile-First Design).",
            "- Reliability: The system should aim for 99.9% uptime during active election periods."
        ]),
        ("Chapter 5: High-Level Design (HLD)", [
            "VotePulse follows a robust, modern Client-Server architecture designed for strict separation of concerns. This separation allows the frontend and backend to scale independently and be developed in parallel.",
            "1. Client Application (Frontend): A Single Page Application (SPA) built with React and Vite. It serves as the presentation layer, handling user interactions, form validations, and routing. It communicates asynchronously with the backend via RESTful APIs using fetch or Axios.",
            "2. API Gateway/Backend Server: A Node.js and Express.js server that acts as the central brain of VotePulse. It processes complex business logic, handles secure authentication flows, validates incoming requests, and serves as the bridge between the client and the persistent storage layer.",
            "3. Database Server: A MySQL relational database. It provides persistent, structured storage with strict ACID properties to guarantee that vote transactions are processed reliably.",
            "4. External Cloud Services:",
            "    - Cloudinary: Utilized as a Content Delivery Network (CDN) and storage solution for candidate and user profile images, offloading storage overhead from the main server.",
            "    - Nodemailer/SMTP: Handles outbound communication, specifically sending verification links and security alerts to user email addresses."
        ]),
        ("Chapter 6: Low-Level Design (LLD)", [
            "The relational database schema is heavily normalized to prevent data redundancy and ensure integrity. It consists of four primary entities mapping to SQL tables:",
            "\nUSER Table:\n- id (Primary Key, UUID)\n- name (String, Not Null)\n- email (String, Unique, Not Null)\n- password (String, Hashed)\n- role (Enum: 'admin', 'voter')\n- isVerified (Boolean, Default: False)",
            "\nELECTION Table:\n- id (Primary Key, UUID)\n- title (String, Not Null)\n- description (Text)\n- startDate (DateTime)\n- endDate (DateTime)\n- status (Enum: 'upcoming', 'active', 'completed')",
            "\nCANDIDATE Table:\n- id (Primary Key, UUID)\n- election_id (Foreign Key referencing Election.id)\n- name (String, Not Null)\n- party (String)\n- image_url (String)",
            "\nVOTE Table (Associative Entity):\n- id (Primary Key, UUID)\n- voter_id (Foreign Key referencing User.id)\n- candidate_id (Foreign Key referencing Candidate.id)\n- election_id (Foreign Key referencing Election.id)\n\nCritical Constraint: A unique composite index on (voter_id, election_id) guarantees that a User can cast only one Vote per Election at the database level. If a duplicate vote is attempted, the database rejects it with an integrity constraint violation, providing bulletproof protection against double-voting."
        ]),
        ("Chapter 7: Tech Stack Details", [
            "Frontend Technologies:",
            "Core Library: React 19. Chosen for its robust component-based architecture, virtual DOM for blazing-fast rendering, and vast ecosystem.",
            "Build Tool: Vite. Chosen for its incredibly fast Hot Module Replacement (HMR) and optimized build process compared to older bundlers like Webpack.",
            "Styling Engine: Tailwind CSS. Allows for highly customizable, utility-first styling without ever leaving the HTML/JSX. Material-UI provides pre-built, accessible components to accelerate development for complex inputs like date-pickers.",
            "Iconography: Lucide-React. Offers a clean, modern, and consistent icon set.",
            "Routing: React Router DOM v7. Industry standard for SPA routing, ensuring URLs update correctly without reloading the entire application.",
            "\nBackend Technologies:",
            "Runtime Environment: Node.js. Allows for a unified JavaScript ecosystem across both the client and server.",
            "Web Framework: Express.js. A lightweight, unopinionated, and highly flexible routing framework.",
            "Database Management: MySQL. A proven, highly reliable relational database system capable of handling complex joins and transactions.",
            "ORM: Sequelize. Simplifies database interactions and schema definitions. It abstracts raw SQL into JavaScript promises, reducing the risk of SQL injection.",
            "Security & Authentication:",
            "- jsonwebtoken (JWT) for stateless session management.",
            "- bcrypt for salting and hashing passwords before storage.",
            "Utilities:",
            "- multer: Middleware for handling multipart/form-data (file uploads).",
            "- cloudinary: SDK for interacting with the Cloudinary media delivery network.",
            "- nodemailer: Module to facilitate sending emails via SMTP."
        ]),
        ("Chapter 8: Advantages & Disadvantages", [
            "Advantages:",
            "- Unparalleled Accessibility: VotePulse democratizes the voting process by allowing users to participate from their homes, workplaces, or while traveling, requiring only an internet connection.",
            "- Massive Cost Reduction: By eliminating the need for physical polling stations, paper ballots, transportation, and manual counting staff, VotePulse offers a highly cost-effective alternative.",
            "- Speed and Accuracy: Vote tallying is instantaneous. Results can be declared the moment an election concludes, completely eliminating human error in counting.",
            "- Environmentally Friendly: The system supports green initiatives by offering a 100% paperless election process.",
            "- Enhanced Auditability: Digital logs and database records provide a clear, auditable trail of election activities (without compromising voter anonymity).",
            "\nDisadvantages:",
            "- The Digital Divide: The system inherently excludes individuals who lack access to smart devices, reliable internet connections, or basic digital literacy.",
            "- Cybersecurity Threats: As a centralized digital platform, VotePulse could be targeted by Distributed Denial of Service (DDoS) attacks, Man-in-the-Middle (MITM) attacks, or sophisticated database breaches.",
            "- Infrastructure Dependency: The entire election process is dependent on server uptime and network stability. A server crash during voting hours could disrupt the election."
        ]),
        ("Chapter 9: API Design and Endpoints", [
            "The following details the primary RESTful API endpoints.",
            "1. POST /api/users/register : Registers a new user. Expects name, email, password. Returns success message and sends verification email.",
            "2. POST /api/users/login : Authenticates user. Expects email, password. Returns JWT token and user profile.",
            "3. POST /api/elections : Creates new election (Admin only). Expects title, startDate, endDate.",
            "4. POST /api/candidates : Registers candidate (Admin only). Uses FormData for name, party, election_id, and image file upload.",
            "5. POST /api/votes/cast : Casts vote (Voter). Expects electionId, candidateId. Requires JWT Authorization Header.",
            "6. GET /api/results/:electionId : Fetches real-time vote aggregations for chart display."
        ]),
        ("Chapter 10: Features & User Journey Results", [
            "The Voter Journey:",
            "1. Registration & Verification: A new user signs up. The system encrypts their password and generates a unique verification token, sending it via Nodemailer to their email. The user clicks the link, verifying their account.",
            "2. Secure Authentication: The user logs in. The backend verifies the Bcrypt hash and issues a JWT, which the frontend stores to authorize subsequent requests.",
            "3. Dashboard Access: The voter lands on their dashboard, viewing active elections tailored to their eligibility.",
            "4. Casting a Vote: The voter navigates to the CastVote interface. They review candidate profiles (images served via Cloudinary), make their selection, and submit.",
            "5. Validation: The backend intercepts the request, verifies the JWT, checks the database to ensure the user hasn't already voted in this election, and then securely records the vote.",
            "\nThe Administrator Journey:",
            "1. Election Creation: An admin accesses the AddElection portal to define the parameters of a new electoral event.",
            "2. Candidate Registration: Using the ManageCandidates portal, the admin registers participants, uploads their photos, and assigns them to the newly created election.",
            "3. User Management: Admins can oversee the user base via ManageUsers, monitoring verification statuses and resolving account issues.",
            "4. Result Monitoring: Once an election concludes, the system aggregates the data, allowing admins to instantly view the final tally and declare winners through interactive charts and tables."
        ]),
        ("Chapter 11: Conclusion & Future Scope", [
            "Conclusion:",
            "VotePulse successfully demonstrates how modern web technologies can be seamlessly leveraged to create a highly secure, scalable, and user-friendly platform for digital elections. By entirely digitizing the voting process, the project addresses and resolves the core inefficiencies of traditional voting—namely cost, accessibility, and speed. The implementation of robust security measures like JWT authentication and relational constraints ensures that the integrity of the electoral process remains intact, paving the way for higher participation rates and transparent, instantaneous results.",
            "\nFuture Scope:",
            "- Blockchain Integration: Transitioning the voting ledger from a centralized MySQL database to a decentralized Blockchain network (using smart contracts). This would make the voting record mathematically immutable and completely immune to database tampering.",
            "- Advanced Biometric Authentication: Integrating with mobile device hardware (FaceID, Fingerprint Scanners) or implementing third-party facial recognition APIs to provide a secondary, highly secure layer of identity verification before a vote is cast.",
            "- AI-Driven Anomaly Detection: Implementing machine learning algorithms to monitor voting patterns in real-time, instantly flagging suspicious activities such as sudden spikes in votes from a single IP address (indicative of bot activity).",
            "- Comprehensive Localization (i18n): Adding multi-language support to ensure the platform is accessible to a diverse, global demographic of voters.",
            "- Advanced Analytics Dashboard: Providing administrators with deep analytical insights, such as voter turnout demographics, peak voting hours, and historical trend comparisons."
        ])
    ]

    for title, paras in chapters:
        add_heading(doc, title, 2)
        for p in paras:
            add_paragraph(doc, p)
        doc.add_page_break()

    # --- IMPLEMENTATION CODE (To increase length to 30-40 pages) ---
    add_heading(doc, "Chapter 12: Implementation Source Code", 1)
    add_paragraph(doc, "This chapter contains the core implementation files of the system, demonstrating the logic for models, controllers, and React components.")
    
    code_files = [
        "../backend/src/models/index.js",
        "../backend/src/modules/users/user.model.js",
        "../backend/src/modules/elections/election.model.js",
        "../backend/src/modules/candidates/candidate.model.js",
        "../backend/src/modules/votes/vote.model.js",
        "../backend/src/modules/auth/auth.controller.js",
        "../backend/src/modules/votes/vote.controller.js",
        "../frontend/src/App.jsx",
        "../frontend/src/pages/CastVote.jsx",
        "../frontend/src/pages/Dashboard.jsx",
        "../frontend/src/pages/ManageCandidates.jsx"
    ]

    for file_path in code_files:
        try:
            if os.path.exists(file_path):
                add_heading(doc, f"File: {os.path.basename(file_path)}", 3)
                with open(file_path, 'r', encoding='utf-8') as f:
                    code_content = f.read()
                
                # Split large code blocks to avoid massive paragraphs
                lines = code_content.split('\n')
                chunk = []
                for line in lines:
                    chunk.append(line)
                    if len(chunk) > 30:
                        p = doc.add_paragraph('\n'.join(chunk))
                        p.runs[0].font.name = 'Courier New'
                        p.runs[0].font.size = Pt(9)
                        chunk = []
                if chunk:
                    p = doc.add_paragraph('\n'.join(chunk))
                    p.runs[0].font.name = 'Courier New'
                    p.runs[0].font.size = Pt(9)
        except Exception as e:
            pass

    # Save document
    doc.save("VotePulse_Final_Report.docx")
    print("Document VotePulse_Final_Report.docx created successfully.")

if __name__ == "__main__":
    create_report()
