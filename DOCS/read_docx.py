import docx
import sys

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        print(f"\n--- Contents of {file_path} (First 2000 chars) ---")
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(f"[{para.style.name}] {para.text}")
        
        text = "\n".join(full_text)
        print(text[:2000])
        print(f"\n--- Total Paragraphs: {len(doc.paragraphs)} ---")
    except Exception as e:
        print(f"Error reading {file_path}: {e}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        read_docx(sys.argv[1])
    else:
        read_docx("sample_report.docx")
        read_docx("old_report.docx")
